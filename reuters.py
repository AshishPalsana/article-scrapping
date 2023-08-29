import os

from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
import time
from selenium.webdriver.support.ui import WebDriverWait

def getarticles(outputpath):
    options = Options()
    driver = webdriver.Chrome(options)

    driver.get('https://www.reuters.com/tags/mergers-acquisitions/')

    links = []

    while True:
        li_objs = driver.find_elements(By.CLASS_NAME, "search-results__item__2oqiX")

        for i in li_objs:
            j = i.find_elements(By.TAG_NAME, "a")
            links.append(j[1].get_attribute("href"))

        break

        # try:
        #     pagination = driver.find_element(By.CLASS_NAME, "search-results__pagination__2h60k")
        #     btns = pagination.find_elements(By.TAG_NAME, "button")
        #
        #     # Check if there's a next button and click it
        #     next_button = None
        #     for btn in btns:
        #         if btn.text.strip().isdigit():
        #             next_button = btn
        #             break
        #
        #     if next_button:
        #         next_button.click()
        #         time.sleep(2)  # Consider using WebDriverWait instead of sleep
        #     else:
        #         break
        # except:
        #     break

    print("articles",len(links))

    if not os.path.exists(outputpath):
        # Create the directory
        os.makedirs(outputpath)

    i = 1
    for link in links:
        driver.get(link)
        title = driver.find_element(By.XPATH,"//header/div/div/h1")
        title = title.text
        # print(title.text)
        body = driver.find_element(By.CLASS_NAME,"article-body__content__17Yit")
        body = body.text
        body = body.replace(r"\nOur Standards: The Thomson Reuters Trust Principles.\nAcquire Licensing Rights\n, opens new tab","")
        # print(body.text)

        doc = Document()

        title_text = f"Title : {title}\n"
        doc.add_paragraph(title_text)

        # doc.add_paragraph()

        doc.add_paragraph(f"Link : {link}")

        # doc.add_paragraph()

        doc.add_paragraph(body)

        path = os.path.join(outputpath, f'{title}.docx')

        doc.save(path)

        if i > 10:
            break
        i+=1
    print("done")

getarticles(r"./output/rereuters")
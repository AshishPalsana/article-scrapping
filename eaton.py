import os.path
import time
import traceback

from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By

def getarticles(outputpath):
    try:
        options = Options()
        options.add_experimental_option("detach", True)
        # options.add_argument("--headless")
        driver = webdriver.Chrome(options)

        driver.delete_all_cookies()

        driver.get("https://www.eaton.com/us/en-us/company/news-insights/news-releases.html")

        while True:
            try:
                driver.find_element(By.XPATH,"//button[text()='Load more']").click()
            except:
                break

        results = driver.find_element(By.CLASS_NAME,'results-list__content')

        links_objs = results.find_elements(By.TAG_NAME,"a")

        links = [obj.get_attribute("href") for obj in links_objs]

        unique_links = list(set(links))

        print("Number of Results :",len(unique_links))

        if not os.path.exists(outputpath):
            # Create the directory
            os.makedirs(outputpath)

        i = 1
        for link in unique_links:
            try:
                time.sleep(2)
                driver.get(link)
                article = driver.find_element(By.XPATH, "html/body/div/div/div/div/div")
                article = article.text
                split_text = article.split('\n', 1)

                doc = Document()

                title = split_text[0]

                title_text = f"Title : {title}\n"
                doc.add_paragraph(title_text)

                # doc.add_paragraph()

                doc.add_paragraph(f"Link : {link}")

                # doc.add_paragraph()

                main_text = split_text[1]

                doc.add_paragraph(main_text)

                path = os.path.join(outputpath, f'{title}.docx')

                doc.save(path)
                if i > 10:
                    break
                i += 1
            except:
                print(link)
                print(traceback.print_exc())
        print("Done")

    except:
        print(traceback.print_exc())

getarticles(r"./output/eaton")